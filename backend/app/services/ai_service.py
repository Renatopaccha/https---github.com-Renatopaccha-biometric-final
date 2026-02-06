"""
AI Service for Google Gemini integration.
Handles chat interactions with session-aware context injection and multimodal file processing.
"""

import asyncio
import base64
import io
from typing import List, Optional, Tuple
import pandas as pd
from PIL import Image
from docx import Document
import google.generativeai as genai

try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("[WARN] PyPDF2 not installed. PDF processing disabled.")

from app.core.config import settings
from app.core.errors import BiometricException
from app.internal.data_manager import data_manager
from app.schemas.ai import FileAttachment


class AIService:
    """
    Service for AI-powered biostatistical assistance using Google Gemini.
    
    Features:
    - Session-aware context injection from active DataFrames
    - Multimodal file processing (images, Word, Excel)
    - Biostatistics tutor persona
    """
    
    def __init__(self):
        """Initialize Gemini client with API key and safety settings."""
        if not settings.gemini_api_key:
            raise BiometricException(
                message="GEMINI_API_KEY not configured in environment variables",
                status_code=500
            )
        
        genai.configure(api_key=settings.gemini_api_key)
        
        # Configure safety settings for scientific/medical content
        self.safety_settings = [
            {
                "category": "HARM_CATEGORY_HARASSMENT",
                "threshold": "BLOCK_NONE"
            },
            {
                "category": "HARM_CATEGORY_HATE_SPEECH",
                "threshold": "BLOCK_NONE"
            },
            {
                "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                "threshold": "BLOCK_NONE"
            },
            {
                "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                "threshold": "BLOCK_NONE"
            }
        ]
        
        self.model = genai.GenerativeModel(
            model_name=settings.gemini_model,
            safety_settings=self.safety_settings
        )
    
    def _build_system_prompt(self, session_context: Optional[str] = None) -> str:
        """
        Build dynamic system prompt with biostatistics expertise and optional session context.
        
        Args:
            session_context: DataFrame metadata and preview from current session
            
        Returns:
            Complete system prompt for the AI
        """
        base_prompt = """Eres un Consultor Experto en Bioestadística y Ciencia de Datos Clínicos para la app "Biometric".
Tu objetivo no es solo responder, sino guiar al investigador (usuario) hacia descubrimientos más profundos.

DIRECTRICES DE RESPUESTA:
1. **Profundidad Analítica:** Nunca des respuestas de una sola línea. Explica el "por qué" estadístico detrás de tu análisis.
2. **Método Socrático:** Al final de cada respuesta importante, DEBES formular una pregunta de seguimiento relevante al usuario. Ejemplo: "¿Consideraste normalizar esta variable antes del test?" o "¿Cuál es tu hipótesis nula para este cruce de datos?".
3. **Contexto:** Si te dan una tabla resumen, analiza tendencias, outliers y valores p. No solo los listes.
4. **Formato:** Usa Markdown para negritas en hallazgos clave y listas para facilitar la lectura.
5. **Elección de Pruebas:** Explica los supuestos de cada prueba estadística y advierte sobre errores comunes (p-hacking, correlación vs causalidad).
6. **Recomendaciones Prácticas:** Sugiere visualizaciones apropiadas y métodos de limpieza de datos cuando sea relevante.
7. **Código:** Formatea código estadístico en bloques cuando sea apropiado (Python/Pandas/SciPy).

Si el usuario pregunta algo simple, responde y luego invítalo a explorar un aspecto más complejo relacionado.

"""
        
        if session_context:
            base_prompt += f"\n**CONTEXTO DE LA SESIÓN ACTUAL:**\n{session_context}\n\n"
            base_prompt += "**IMPORTANTE:** El usuario tiene estos datos cargados. Cuando hagas sugerencias, menciona las columnas específicas que están disponibles. No inventes nombres de columnas que no existen.\n\n"
        
        return base_prompt
    
    def _get_session_context(self, session_id: str) -> Optional[str]:
        """
        Retrieve DataFrame metadata and preview from session.
        
        Args:
            session_id: Active session identifier
            
        Returns:
            Formatted string with DataFrame information or None if session invalid
        """
        try:
            df = data_manager.get_dataframe(session_id)
            
            # Build context string
            context_parts = [
                f"📊 **Dataset cargado:** {len(df)} filas × {len(df.columns)} columnas\n",
                f"**Columnas disponibles:**"
            ]

            # Vectorized null count calculation (performance optimization)
            null_counts = df.isnull().sum()
            null_pcts = (null_counts / len(df) * 100)

            # List columns with data types
            for col in df.columns:
                dtype = df[col].dtype
                null_pct = null_pcts[col]

                # Classify data type for user
                if pd.api.types.is_numeric_dtype(df[col]):
                    type_label = "numérica"
                elif pd.api.types.is_datetime64_any_dtype(df[col]):
                    type_label = "fecha/hora"
                else:
                    type_label = "categórica/texto"

                context_parts.append(
                    f"  - `{col}` ({type_label}, {dtype}) - {null_pct:.1f}% nulos"
                )
            
            # Add data preview (first 5 rows)
            context_parts.append(f"\n**Vista previa (primeras 5 filas):**")
            context_parts.append("```")
            context_parts.append(df.head(5).to_string())
            context_parts.append("```")
            
            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
            if numeric_cols:
                context_parts.append(f"\n**Estadísticas descriptivas (columnas numéricas):**")
                context_parts.append("```")
                context_parts.append(df[numeric_cols].describe().to_string())
                context_parts.append("```")
            
            return "\n".join(context_parts)
            
        except Exception as e:
            # Session might be expired or invalid
            return None
    
    def _decode_attachment(self, attachment: FileAttachment) -> Tuple[bytes, str]:
        """
        Decode Base64 attachment with robust handling of data URI prefixes.
        
        Args:
            attachment: FileAttachment with Base64 data
            
        Returns:
            Tuple of (file_bytes, filename)
        """
        try:
            # Clean Base64 data - remove data URI prefix if present
            # Format: "data:mime/type;base64,ACTUAL_BASE64_DATA"
            base64_data = attachment.data
            
            if "base64," in base64_data:
                # Extract only the Base64 part after the comma
                base64_data = base64_data.split("base64,")[1]
            elif "," in base64_data and base64_data.startswith("data:"):
                # Handle other data URI formats
                base64_data = base64_data.split(",", 1)[1]
            
            # Decode Base64 to bytes
            file_bytes = base64.b64decode(base64_data)
            
            print(f"[DEBUG] Successfully decoded attachment '{attachment.name}', size: {len(file_bytes)} bytes")
            
            return file_bytes, attachment.name
            
        except Exception as e:
            error_msg = f"Failed to decode attachment '{attachment.name}': {str(e)}"
            print(f"[ERROR] {error_msg}")
            print(f"[DEBUG] Base64 data preview: {attachment.data[:100]}...")
            raise BiometricException(
                message=error_msg,
                status_code=400
            )
    
    def _validate_file(self, file_content: bytes, filename: str) -> None:
        """
        Validate file size and type.
        
        Args:
            file_content: File bytes
            filename: Original filename
            
        Raises:
            BiometricException: If file is invalid
        """
        # Check size
        size_mb = len(file_content) / (1024 * 1024)
        if size_mb > settings.max_ai_file_size_mb:
            raise BiometricException(
                message=f"File '{filename}' exceeds maximum size of {settings.max_ai_file_size_mb}MB",
                status_code=400
            )
        
        # Check extension
        allowed_extensions = ['.png', '.jpg', '.jpeg', '.docx', '.xlsx', '.xls', '.pdf', '.csv']
        if not any(filename.lower().endswith(ext) for ext in allowed_extensions):
            raise BiometricException(
                message=f"File type not supported. Allowed: {', '.join(allowed_extensions)}",
                status_code=400
            )
    
    def _process_image(self, file_content: bytes, filename: str) -> Tuple[Image.Image, str]:
        """
        Process image file for Gemini API.
        
        Args:
            file_content: Image bytes
            filename: Original filename
            
        Returns:
            Tuple of (PIL Image, description text)
        """
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Convert RGBA to RGB if necessary
            if image.mode == 'RGBA':
                image = image.convert('RGB')
            
            description = f"📷 Imagen adjunta: {filename}"
            return image, description
            
        except Exception as e:
            raise BiometricException(
                message=f"Failed to process image '{filename}': {str(e)}",
                status_code=400
            )
    
    def _process_docx(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from Word document.
        
        Args:
            file_content: Document bytes
            filename: Original filename
            
        Returns:
            Extracted text with formatting
        """
        try:
            doc = Document(io.BytesIO(file_content))
            
            text_parts = [f"📄 **Documento Word adjunto: {filename}**\n"]
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                text_parts.append("\n[Tabla detectada]")
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            raise BiometricException(
                message=f"Failed to process Word document '{filename}': {str(e)}",
                status_code=400
            )
    
    def _process_excel_sync(self, file_content: bytes, filename: str) -> str:
        """
        Synchronous Excel file processing (runs in thread pool).

        Args:
            file_content: Excel bytes
            filename: Original filename

        Returns:
            Markdown-formatted table (limited to first 50 rows)
        """
        try:
            print(f"[DEBUG] Processing Excel file: {filename}, size: {len(file_content)} bytes")

            # Read Excel file
            df = pd.read_excel(io.BytesIO(file_content))

            print(f"[DEBUG] Excel loaded: {len(df)} rows, {len(df.columns)} columns")

            if df.empty:
                return f"📊 **Archivo Excel adjunto: {filename}**\n\n⚠️ El archivo está vacío."

            # Limit rows for performance
            max_rows = 50
            if len(df) > max_rows:
                df_preview = df.head(max_rows)
                truncation_note = f"\n\n*(Mostrando primeras {max_rows} filas de {len(df)} totales)*"
            else:
                df_preview = df
                truncation_note = ""

            # Convert to markdown
            markdown = f"📊 **Archivo Excel adjunto: {filename}**\n\n"
            markdown += f"**Dimensiones:** {len(df)} filas × {len(df.columns)} columnas\n\n"
            markdown += df_preview.to_markdown(index=False)
            markdown += truncation_note

            return markdown

        except PermissionError:
            error_msg = f"❌ No se pudo leer el archivo Excel '{filename}': archivo protegido con contraseña o sin permisos."
            print(f"[ERROR] {error_msg}")
            raise BiometricException(message=error_msg, status_code=400)
        except pd.errors.EmptyDataError:
            error_msg = f"❌ El archivo Excel '{filename}' está vacío o corrupto."
            print(f"[ERROR] {error_msg}")
            raise BiometricException(message=error_msg, status_code=400)
        except Exception as e:
            error_msg = f"❌ Error al procesar el archivo Excel '{filename}': {str(e)}"
            print(f"[ERROR] {error_msg}")
            raise BiometricException(message=error_msg, status_code=400)

    async def _process_excel(self, file_content: bytes, filename: str) -> str:
        """
        Read Excel file and convert to markdown table (async).
        Runs blocking I/O operations in a thread pool to avoid blocking the event loop.

        Args:
            file_content: Excel bytes
            filename: Original filename

        Returns:
            Markdown-formatted table (limited to first 50 rows)
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self._process_excel_sync,
            file_content,
            filename
        )
    
    def _process_csv_sync(self, file_content: bytes, filename: str) -> str:
        """
        Synchronous CSV file processing (runs in thread pool).

        Args:
            file_content: CSV bytes
            filename: Original filename

        Returns:
            Markdown-formatted table (limited to first 50 rows)
        """
        try:
            print(f"[DEBUG] Processing CSV file: {filename}, size: {len(file_content)} bytes")

            # Read CSV file
            df = pd.read_csv(io.BytesIO(file_content))

            print(f"[DEBUG] CSV loaded: {len(df)} rows, {len(df.columns)} columns")

            if df.empty:
                return f"📊 **Archivo CSV adjunto: {filename}**\n\n⚠️ El archivo está vacío."

            # Limit rows for performance
            max_rows = 50
            if len(df) > max_rows:
                df_preview = df.head(max_rows)
                truncation_note = f"\n\n*(Mostrando primeras {max_rows} filas de {len(df)} totales)*"
            else:
                df_preview = df
                truncation_note = ""

            # Convert to markdown
            markdown = f"📊 **Archivo CSV adjunto: {filename}**\n\n"
            markdown += f"**Dimensiones:** {len(df)} filas × {len(df.columns)} columnas\n\n"
            markdown += df_preview.to_markdown(index=False)
            markdown += truncation_note

            return markdown

        except pd.errors.EmptyDataError:
            error_msg = f"❌ El archivo CSV '{filename}' está vacío."
            print(f"[ERROR] {error_msg}")
            raise BiometricException(message=error_msg, status_code=400)
        except pd.errors.ParserError as e:
            error_msg = f"❌ Error al parsear el archivo CSV '{filename}': formato inválido. {str(e)}"
            print(f"[ERROR] {error_msg}")
            raise BiometricException(message=error_msg, status_code=400)
        except Exception as e:
            error_msg = f"❌ Error al procesar el archivo CSV '{filename}': {str(e)}"
            print(f"[ERROR] {error_msg}")
            raise BiometricException(message=error_msg, status_code=400)

    async def _process_csv(self, file_content: bytes, filename: str) -> str:
        """
        Read CSV file and convert to markdown table (async).
        Runs blocking I/O operations in a thread pool to avoid blocking the event loop.

        Args:
            file_content: CSV bytes
            filename: Original filename

        Returns:
            Markdown-formatted table (limited to first 50 rows)
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self._process_csv_sync,
            file_content,
            filename
        )
    
    def _process_pdf_sync(self, file_content: bytes, filename: str) -> str:
        """
        Synchronous PDF file processing (runs in thread pool).

        Args:
            file_content: PDF bytes
            filename: Original filename

        Returns:
            Extracted text with formatting
        """
        if not PDF_SUPPORT:
            raise BiometricException(
                message="❌ Procesamiento de PDF no disponible. La librería PyPDF2 no está instalada.",
                status_code=400
            )

        try:
            print(f"[DEBUG] Processing PDF file: {filename}, size: {len(file_content)} bytes")

            pdf_reader = PdfReader(io.BytesIO(file_content))

            print(f"[DEBUG] PDF loaded: {len(pdf_reader.pages)} pages")

            text_parts = [f"📄 **Archivo PDF adjunto: {filename}**\n"]
            text_parts.append(f"**Páginas:** {len(pdf_reader.pages)}\n\n")

            # Extract text from all pages (limit to first 10 for performance)
            max_pages = 10
            pages_with_text = 0
            for i, page in enumerate(pdf_reader.pages[:max_pages]):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"**Página {i+1}:**\n{page_text}\n\n")
                    pages_with_text += 1

            if pages_with_text == 0:
                text_parts.append("⚠️ No se pudo extraer texto del PDF. Puede ser un PDF escaneado (imagen) o protegido.")

            if len(pdf_reader.pages) > max_pages:
                text_parts.append(f"\n*(Mostrando primeras {max_pages} páginas de {len(pdf_reader.pages)} totales)*")

            return "\n".join(text_parts)

        except Exception as e:
            error_msg = f"❌ Error al procesar el archivo PDF '{filename}': {str(e)}"
            print(f"[ERROR] {error_msg}")
            raise BiometricException(message=error_msg, status_code=400)

    async def _process_pdf(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from PDF file (async).
        Runs blocking I/O operations in a thread pool to avoid blocking the event loop.

        Args:
            file_content: PDF bytes
            filename: Original filename

        Returns:
            Extracted text with formatting
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self._process_pdf_sync,
            file_content,
            filename
        )
    
    async def chat(
        self,
        message: str,
        session_id: Optional[str] = None,
        history: List[dict] = None,
        attachments: List[FileAttachment] = None
    ) -> dict:
        """
        Generate AI response with optional session context and file attachments.
        
        Args:
            message: User's message
            session_id: Optional session ID for DataFrame context
            history: Previous conversation messages
            attachments: List of FileAttachment objects with Base64 data
            
        Returns:
            Dict with response, context_used flag, and files_processed count
        """
        try:
            # Get session context if available
            session_context = None
            if session_id:
                session_context = self._get_session_context(session_id)
            
            # Build system prompt
            system_prompt = self._build_system_prompt(session_context)
            
            # Process uploaded files from Base64 attachments
            file_contents = []
            files_processed = 0
            
            if attachments:
                for attachment in attachments:
                    # Decode Base64
                    file_content, filename = self._decode_attachment(attachment)
                    
                    # Validate file
                    self._validate_file(file_content, filename)
                    
                    # Process based on type
                    if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                        image, description = self._process_image(file_content, filename)
                        file_contents.append(image)
                        message += f"\n\n{description}"
                        files_processed += 1
                        
                    elif filename.lower().endswith('.docx'):
                        text = self._process_docx(file_content, filename)
                        message += f"\n\n{text}"
                        files_processed += 1
                        
                    elif filename.lower().endswith(('.xlsx', '.xls')):
                        table = await self._process_excel(file_content, filename)
                        message += f"\n\n{table}"
                        files_processed += 1

                    elif filename.lower().endswith('.csv'):
                        table = await self._process_csv(file_content, filename)
                        message += f"\n\n{table}"
                        files_processed += 1

                    elif filename.lower().endswith('.pdf'):
                        text = await self._process_pdf(file_content, filename)
                        message += f"\n\n{text}"
                        files_processed += 1
            
            # Build conversation history for context
            conversation_parts = [system_prompt]
            
            if history:
                for msg in history:
                    # Access as Pydantic object attributes (not dict)
                    role = msg.role if hasattr(msg, 'role') else msg.get("role", "user")
                    content = msg.content if hasattr(msg, 'content') else msg.get("content", "")
                    
                    if role == "user":
                        conversation_parts.append(f"Usuario: {content}")
                    else:
                        conversation_parts.append(f"Asistente: {content}")
            
            # Add current user message
            conversation_parts.append(f"Usuario: {message}")
            
            # Combine text and images
            full_prompt = "\n\n".join(conversation_parts)
            
            # Generate response
            if file_contents:
                # Multimodal request (text + images)
                response = self.model.generate_content(
                    [full_prompt] + file_contents,
                    generation_config={
                        "temperature": 0.7,
                        "top_p": 0.95,
                        "top_k": 40,
                        "max_output_tokens": 4096,
                    }
                )
            else:
                # Text-only request
                response = self.model.generate_content(
                    full_prompt,
                    generation_config={
                        "temperature": 0.7,
                        "top_p": 0.95,
                        "top_k": 40,
                        "max_output_tokens": 4096,
                    }
                )
            
            return {
                "success": True,
                "response": response.text,
                "session_context_used": session_context is not None,
                "files_processed": files_processed,
                "error": None
            }
            
        except BiometricException:
            # Re-raise our custom exceptions
            raise
            
        except Exception as e:
            # Handle any other errors gracefully
            error_msg = str(e)
            
            # Check for common API errors
            if "API_KEY" in error_msg.upper():
                error_msg = "API key inválida o no configurada. Verifica GEMINI_API_KEY en variables de entorno."
            elif "QUOTA" in error_msg.upper() or "RATE_LIMIT" in error_msg.upper():
                error_msg = "Límite de uso de la API alcanzado. Intenta nuevamente en unos minutos."
            
            return {
                "success": False,
                "response": "Lo siento, ocurrió un error al procesar tu solicitud. Por favor, intenta nuevamente.",
                "session_context_used": False,
                "files_processed": 0,
                "error": error_msg
            }

    async def interpret_statistics(self, stats_data: dict, segment_name: str = "General") -> str:
        """
        Genera una interpretación narrativa profesional de la tabla de estadísticas.
        
        Args:
            stats_data: Dictionary with variable names as keys and their statistics as values
            segment_name: Name of the segment being analyzed (default: "General")
            
        Returns:
            Markdown-formatted interpretation string
        """
        import json
        
        # 1. Definimos la PERSONA y el OBJETIVO CLARO
        system_instruction = """
        Actúa como un Bioestadístico Senior y Consultor de Investigación Clínica para la plataforma 'Biometric'.
        Tu objetivo es analizar una tabla de estadísticas descriptivas y generar un resumen interpretativo 
        que el investigador pueda copiar y pegar directamente en la sección de "Resultados preliminares" de su paper o tesis.
        """

        # 2. Definimos las REGLAS DE ANÁLISIS (El "Cerebro" del prompt)
        analysis_rules = """
        AL ANALIZAR LOS DATOS, BUSCA ACTIVAMENTE:
        1. **Prueba de Normalidad (Crucial):** Revisa el campo 'normality_test'. 
           - Si dice 'Normal', confirma que se pueden usar pruebas paramétricas (t-Student, ANOVA).
           - Si dice 'No Normal', SUGIERE OBLIGATORIAMENTE usar pruebas no paramétricas (U de Mann-Whitney, Kruskal-Wallis).
           
        2. **Homogeneidad de los datos:** Mira el Coeficiente de Variación (cv).
           - Si es < 10%: Datos muy homogéneos.
           - Si es > 30%: Datos muy dispersos/heterogéneos (posibles valores atípicos o muestras mixtas).
           
        3. **Asimetría:** Compara la Media vs. Mediana.
           - Si difieren significativamente (>10%), menciona hacia dónde está el sesgo (asimetría positiva/negativa).
        
        4. **Hallazgos Clínicos (Contexto):** Usa los nombres de las variables.
           - Ejemplo: Si la variable es "Edad" y la media es 85, comenta que es una población geriátrica. 
           - Si es "Glucosa" y la media es 200, alerta sobre posibles valores diabéticos en la muestra.
        """

        # 3. Definimos el FORMATO DE SALIDA (Para que se vea bonito en el Frontend)
        output_format = f"""
        FORMATO DE RESPUESTA (Usa Markdown):
        
        ### 🧠 Interpretación Automática: Segmento '{segment_name}'
        
        **Resumen General:**
        [Un párrafo de 3-4 líneas resumiendo el comportamiento general de las variables]
        
        **Hallazgos Clave:**
        * ✅ **Distribución:** [Comentario sobre Normalidad y qué test usar]
        * 📊 **Dispersión:** [Comentario sobre el CV y estabilidad de datos]
        * 🚨 **Atención:** [Menciona cualquier variable con asimetría fuerte o valores extremos/outliers detectados por la curtosis]
        
        **Recomendación:**
        [Una frase final aconsejando el siguiente paso analítico]
        """

        # Construimos el prompt final
        prompt = f"{system_instruction}\n\n{analysis_rules}\n\n{output_format}\n\nDATOS A ANALIZAR:\n{json.dumps(stats_data, indent=2)}"

        try:
            # Usamos una temperatura baja (0.3 - 0.4) para que sea analítico y preciso, no "creativo"
            response = self.model.generate_content(
                prompt,
                generation_config={"temperature": 0.35} 
            )
            return response.text
        except Exception as e:
            return f"Error al generar interpretación: {str(e)}"


# Singleton instance - only instantiate if API key is configured
ai_service = None
try:
    if settings.gemini_api_key:
        ai_service = AIService()
except Exception as e:
    # AI service not available, but don't crash the server
    print(f"Warning: AI Service not initialized: {e}")
    ai_service = None

